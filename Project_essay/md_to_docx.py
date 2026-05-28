"""Converte Artigo_SBC_VR_Folklore.md -> .docx em formato aproximado do SBC.

O que o script faz automaticamente:
- usa o titulo escolhido (Opcao 1) + monta o bloco de autores/filiacao
- descarta os andaimes do .md (secao "Sugestoes de titulo" e o aviso "Como usar este arquivo")
- A4, margens SBC-like, corpo em Times New Roman 12 justificado
- EMBUTE as Figuras 2 e 4 (diagramas gerados por gen_figuras.py) com legenda
- cria placeholders marcados para as Figuras 1, 3, 5, 6 e 7 (screenshots do Unity)
- monta a Tabela 1 com legenda

Rode dentro de Project_essay/: python md_to_docx.py
(rode antes: python gen_figuras.py, para os diagramas existirem)
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = Path(__file__).parent
MD_PATH = BASE / "Artigo_SBC_VR_Folklore.md"
DOCX_PATH = BASE / "Artigo_SBC_VR_Folklore.docx"

CENTER = WD_ALIGN_PARAGRAPH.CENTER
JUSTIFY = WD_ALIGN_PARAGRAPH.JUSTIFY

TITLE = ("Explorando Histórias: Desenvolvimento de uma Experiência em Realidade "
         "Virtual com Narrativas Folclóricas Continentais")
AUTHORS = "Leonardo Buratto de Assis, Jennifer Leonora Galina Vieira, Leonardo Conti"
AFFIL = "Centro Universitário Sagrado Coração (Unisagrado) – Bauru – SP – Brasil"

# Figuras 2 e 4 ja existem (geradas por gen_figuras.py). As demais sao screenshots
# do Unity: basta salvar o PNG com o nome abaixo dentro de figuras/ e rodar de novo
# que o script embute automaticamente; enquanto nao existir, fica um placeholder.
FIG_IMAGES = {
    1: BASE / "figuras" / "fig1_lobby.png",
    2: BASE / "figuras" / "fig2_arquitetura.png",
    3: BASE / "figuras" / "fig3_crosshair.png",
    4: BASE / "figuras" / "fig4_maquina_estados.png",
    5: BASE / "figuras" / "fig5_hierarchy.png",
    6: BASE / "figuras" / "fig6_lobby_topdown.png",
    7: BASE / "figuras" / "fig7_npcs.png",
}
TABLE_LABEL = "Tabela 1"
TABLE_CAPTION = "NPCs implementados, suas cenas e os temas de narrativa."


# --------------------------------------------------------------------------- #
# Formatacao inline (negrito / italico / codigo)
# --------------------------------------------------------------------------- #
def parse_inline(text):
    parts = []
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            parts.append((text[pos:m.start()], False, False, False))
        chunk = m.group(0)
        if chunk.startswith("**"):
            parts.append((chunk[2:-2], True, False, False))
        elif chunk.startswith("`"):
            parts.append((chunk[1:-1], False, False, True))
        else:
            parts.append((chunk[1:-1], False, True, False))
        pos = m.end()
    if pos < len(text):
        parts.append((text[pos:], False, False, False))
    return parts


def add_runs(p, text, size=None):
    for txt, bold, italic, is_code in parse_inline(text):
        run = p.add_run(txt)
        run.bold = bold
        run.italic = italic
        if size:
            run.font.size = Pt(size)
        if is_code:
            run.font.name = "Consolas"
            run.font.size = Pt(size - 2 if size else 10)
    return p


def body_paragraph(doc, text, style=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    add_runs(p, text)
    if style is None:
        p.alignment = JUSTIFY
    return p


def section_heading(doc, text, size, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def centered_bold(doc, text, size=12):
    p = doc.add_paragraph()
    p.alignment = CENTER
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    return p


def add_caption(doc, label, text):
    p = doc.add_paragraph()
    p.alignment = CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(label + ". ")
    run.bold = True
    run.font.size = Pt(10)
    add_runs(p, text, size=10)
    return p


def add_figure(doc, num, caption):
    img = FIG_IMAGES.get(num)
    if img and img.exists():
        doc.add_picture(str(img), width=Inches(6.0))
        doc.paragraphs[-1].alignment = CENTER
    else:
        fname = img.name if img else f"fig{num}.png"
        p = doc.add_paragraph()
        p.alignment = CENTER
        p.paragraph_format.space_before = Pt(8)
        run = p.add_run(f"[  Figura {num}: salve o screenshot como  figuras\\{fname}  "
                        f"e rode md_to_docx.py novamente  ]")
        run.italic = True
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    add_caption(doc, f"Figura {num}", caption)


# --------------------------------------------------------------------------- #
# Pre-processamento: remove andaimes e gera marcadores de figura/tabela
# --------------------------------------------------------------------------- #
def filter_lines(raw):
    out = []
    lines = raw.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        s = line.strip()

        # Placeholder de FIGURA (blockquote) -> marcador @@FIG@@num@@legenda
        if s.startswith("> **[FIGURA"):
            buf = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                buf.append(lines[i].strip().lstrip(">").strip())
                i += 1
            joined = " ".join(buf)
            m = re.search(r"\[FIGURA\s+(\d+).*?\]\*\*\s*(.*)", joined)
            if m:
                num = m.group(1)
                cap = re.split(r"\*Sugest", m.group(2))[0]
                cap = re.sub(r"\s+", " ", cap).strip().rstrip(".") + "."
                out.append(f"@@FIG@@{num}@@{cap}")
            if i < len(lines) and lines[i].strip() == "":
                i += 1
            continue

        # Placeholder de TABELA -> marcador @@TABCAP@@ + linhas da tabela
        if s.startswith("> **[TABELA"):
            out.append("@@TABCAP@@")
            i += 1
            while i < len(lines) and lines[i].strip() in (">", ""):
                i += 1
            while i < len(lines) and lines[i].strip().startswith("> |"):
                out.append(lines[i].strip()[2:])
                i += 1
            continue

        out.append(line)
        i += 1
    return "\n".join(out)


# --------------------------------------------------------------------------- #
# Tabela markdown -> tabela Word
# --------------------------------------------------------------------------- #
def build_table(doc, table_buf):
    headers = [c.strip() for c in table_buf[0].strip("|").split("|")]
    data = []
    for row in table_buf[2:]:  # pula a linha separadora |---|
        cells = [c.strip() for c in row.strip("|").split("|")]
        cells = (cells + [""] * len(headers))[:len(headers)]
        data.append(cells)
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        pass
    for idx, h in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
    for r_idx, row in enumerate(data):
        for c_idx, txt in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            add_runs(cell.paragraphs[0], txt)
    doc.add_paragraph()


# --------------------------------------------------------------------------- #
# Conversao principal
# --------------------------------------------------------------------------- #
def md_to_docx():
    raw = MD_PATH.read_text(encoding="utf-8")
    # descarta tudo antes do Abstract (titulo/sugestoes/autores do .md viram bloco proprio)
    if "## Abstract" in raw:
        raw = raw[raw.index("## Abstract"):]
    raw = filter_lines(raw)

    doc = Document()

    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(3.0)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    # ---- bloco de titulo / autores ----
    pt = doc.add_paragraph()
    pt.alignment = CENTER
    rt = pt.add_run(TITLE)
    rt.bold = True
    rt.font.size = Pt(16)

    doc.add_paragraph()
    pa = doc.add_paragraph()
    pa.alignment = CENTER
    pa.add_run(AUTHORS).font.size = Pt(12)
    pf = doc.add_paragraph()
    pf.alignment = CENTER
    rf = pf.add_run(AFFIL)
    rf.font.size = Pt(12)
    rf.italic = True
    doc.add_paragraph()

    # ---- corpo ----
    lines = raw.split("\n")
    i = 0
    in_code = False
    code_buf = []
    table_buf = []

    while i < len(lines):
        line = lines[i]
        s = line.strip()

        # marcadores
        if s.startswith("@@FIG@@"):
            parts = s.split("@@", 3)
            add_figure(doc, int(parts[2]), parts[3])
            i += 1
            continue
        if s == "@@TABCAP@@":
            add_caption(doc, TABLE_LABEL, TABLE_CAPTION)
            i += 1
            continue

        # codigo
        if s.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buf))
                run.font.name = "Consolas"
                run.font.size = Pt(10)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # tabela
        if s.startswith("|") and s.endswith("|"):
            table_buf.append(s)
            i += 1
            continue
        if table_buf:
            build_table(doc, table_buf)
            table_buf = []

        # cabecalhos
        if s.startswith("## "):
            htext = s[3:].strip()
            if htext in ("Abstract", "Resumo"):
                centered_bold(doc, htext, size=12)
            else:
                section_heading(doc, htext, size=13)
        elif s.startswith("### "):
            section_heading(doc, s[4:].strip(), size=12)
        elif s.startswith("#### "):
            section_heading(doc, s[5:].strip(), size=11, italic=True)
        elif s.startswith("- "):
            body_paragraph(doc, s[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s", s):
            body_paragraph(doc, re.sub(r"^\d+\.\s", "", s).strip(), style="List Number")
        elif s == "---" or s == "":
            pass
        elif s.startswith("> "):
            body_paragraph(doc, s[2:].strip())
        else:
            body_paragraph(doc, s)

        i += 1

    if table_buf:
        build_table(doc, table_buf)

    doc.save(DOCX_PATH)
    print(f"OK -> {DOCX_PATH}")


if __name__ == "__main__":
    md_to_docx()
